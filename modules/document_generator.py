# modules/document_generator.py
import os
import markdown
from docx import Document
from datetime import datetime
import re
from docx.shared import RGBColor

from config import OUTPUT_DIR, TEMPLATE_DIR, LAB_BOOK_SECTIONS

class DocumentGenerator:
    def __init__(self):
        """Initialize document generator with default templates"""
        # Create output directory if it doesn't exist
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        # Default template file
        self.template_file = os.path.join(TEMPLATE_DIR, "labbook_template.md")
        
        # Create default template if it doesn't exist
        if not os.path.exists(self.template_file):
            os.makedirs(TEMPLATE_DIR, exist_ok=True)
            self._create_default_template()
    
    def _create_default_template(self):
        """Create a default Markdown template for lab books"""
        template = "# {title}\n\n"
        template += "**Date:** {date}\n\n"
        template += "**Participants:** {participants}\n\n"
        
        for section in LAB_BOOK_SECTIONS[3:]:  # Skip title, date, participants
            section_key = section.lower().replace(' ', '_')
            template += f"## {section}\n\n"
            
            if section == "External Analysis":
                template += "<span style=\"color:red\">\n{content_" + section_key + "}\n</span>\n\n"
            elif section == "External Comments":
                template += "<span style=\"color:blue\">\n{content_" + section_key + "}\n</span>\n\n"
            else:
                template += "{content_" + section_key + "}\n\n"
        
        # Save template
        with open(self.template_file, 'w') as f:
            f.write(template)
        
        print(f"Created default template at {self.template_file}")
    
    def generate_markdown(self, content, title=None):
        """Generate a Markdown document from structured content"""
        # Create a timestamp for the filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title = title or f"Lab Book Entry {timestamp}"
        filename = os.path.join(OUTPUT_DIR, f"labbook_{timestamp}.md")
        
        # Write content to file
        with open(filename, 'w') as f:
            f.write(content)
        
        print(f"Markdown lab book saved to {filename}")
        return filename
    
    def generate_docx(self, content, title=None):
        """Generate a Word document from Markdown content with colored text"""
        # Create a timestamp for the filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title = title or f"Lab Book Entry {timestamp}"
        filename = os.path.join(OUTPUT_DIR, f"labbook_{timestamp}.docx")
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Convert markdown to basic formatting
        sections = content.split('##')
        
        # Add the content before the first section (usually intro, date, etc.)
        if sections[0].strip():
            for para in sections[0].split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Process each section
        for section in sections[1:]:
            if not section.strip():
                continue
                
            # Extract section title and content
            section_lines = section.strip().split('\n')
            section_title = section_lines[0].strip()
            section_content = '\n'.join(section_lines[1:]).strip()
            
            # Add section heading
            doc.add_heading(section_title, 1)
            
            # Check for colored text sections
            is_red_section = "External Analysis" in section_title
            is_blue_section = "External Comments" in section_title
            
            # Remove HTML color span tags if present
            section_content = re.sub(r'<span style="color:red">', '', section_content)
            section_content = re.sub(r'<span style="color:blue">', '', section_content)
            section_content = re.sub(r'</span>', '', section_content)
            
            # Add content with appropriate color
            if section_content:
                for para in section_content.split('\n'):
                    if para.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(para.strip())
                        if is_red_section:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                        elif is_blue_section:
                            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        
        # Save the document
        doc.save(filename)
        
        print(f"Word document lab book saved to {filename}")
        return filename
    
    def add_image_to_document(self, doc_path, image_path, caption=None):
        """Add an image to an existing document"""
        # Check if the document is markdown or docx
        if doc_path.endswith('.md'):
            # Add image to markdown
            with open(doc_path, 'a') as f:
                f.write(f"\n\n![{caption or 'Image'}]({image_path})\n")
                if caption:
                    f.write(f"*{caption}*\n")
            return True
            
        elif doc_path.endswith('.docx'):
            # Add image to docx
            doc = Document(doc_path)
            doc.add_picture(image_path)
            if caption:
                doc.add_paragraph(caption, style='Caption')
            doc.save(doc_path)
            return True
            
        else:
            print(f"Unsupported document format: {doc_path}")
            return False


# Simple command-line test
if __name__ == "__main__":
    generator = DocumentGenerator()
    
    # Test with some sample content
    sample_content = """# Sample Lab Experiment
    
**Date:** 2023-09-20

**Participants:** Dr. Smith, Jane Doe, John Smith

## Aims
To determine the reaction rate of compound X with catalyst Y.

## Choices
- Selected catalyst Y due to its higher surface area
- Used stirring speed of 300 rpm to ensure uniform mixing
- Measured at 5-minute intervals for better resolution

## Summary
The reaction showed first-order kinetics with respect to compound X.
Rate constant was determined to be 0.045 min^-1.

## Questions
- How would temperature affect the reaction rate?
- Would a different solvent change the kinetics?
- What is the mechanism for this reaction?

## External Analysis
<span style="color:red">
This experiment effectively established baseline kinetics for the X+Y reaction system.
The choice of catalyst was appropriate, though testing multiple catalysts would provide valuable comparative data.
Consider running the reaction at 3-4 different temperatures to determine activation energy.
</span>

## External Comments
<span style="color:blue">
[Dr. Smith]: Remember to consider the effect of oxygen on these reactions.
[Jane Doe]: The UV spectra showed an unexpected peak at 340nm we should investigate.
</span>
"""
    
    md_file = generator.generate_markdown(sample_content, "Reaction Rate Study")
    docx_file = generator.generate_docx(sample_content, "Reaction Rate Study")
    
    print(f"Test files created: {md_file} and {docx_file}")