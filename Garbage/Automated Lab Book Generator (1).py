# modules/document_generator.py
import os
import markdown
from docx import Document
from datetime import datetime

from config import OUTPUT_DIR, TEMPLATE_DIR, LAB_BOOK_SECTIONS

class DocumentGenerator:
    def __init__(self):
        """Initialize document generator with default templates"""
        # Create output directory if it doesn't exist
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        # Default template file
        self.template_file = os.path.join(TEMPLATE_DIR, "labbook_template.md")
        
        # Create default template if it doesn't exist
        if not os.path.exists(self.template_file):
            os.makedirs(TEMPLATE_DIR, exist_ok=True)
            self._create_default_template()
    
    def _create_default_template(self):
        """Create a default Markdown template for lab books"""
        template = "# {title}\n\n"
        template += "**Date:** {date}\n\n"
        template += "**Participants:** {participants}\n\n"
        
        for section in LAB_BOOK_SECTIONS[3:]:  # Skip title, date, participants
            template += f"## {section}\n\n{{content_{section.lower().replace(' ', '_')}}}\n\n"
        
        # Save template
        with open(self.template_file, 'w') as f:
            f.write(template)
        
        print(f"Created default template at {self.template_file}")
    
    def generate_markdown(self, content, title=None):
        """Generate a Markdown document from structured content"""
        # Create a timestamp for the filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title = title or f"Lab Book Entry {timestamp}"
        filename = os.path.join(OUTPUT_DIR, f"labbook_{timestamp}.md")
        
        # Write content to file
        with open(filename, 'w') as f:
            f.write(content)
        
        print(f"Markdown lab book saved to {filename}")
        return filename
    
    def generate_docx(self, content, title=None):
        """Generate a Word document from Markdown content"""
        # Create a timestamp for the filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        title = title or f"Lab Book Entry {timestamp}"
        filename = os.path.join(OUTPUT_DIR, f"labbook_{timestamp}.docx")
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Convert markdown to basic formatting
        # This is simplified - a full implementation would parse the markdown properly
        sections = content.split('##')
        
        # Add the content before the first section (usually intro, date, etc.)
        if sections[0].strip():
            for para in sections[0].split('\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
        
        # Process each section
        for section in sections[1:]:
            if not section.strip():
                continue
                
            # Extract section title and content
            section_lines = section.strip().split('\n')
            section_title = section_lines[0].strip()
            section_content = '\n'.join(section_lines[1:]).strip()
            
            # Add section heading and content
            doc.add_heading(section_title, 1)
            if section_content:
                for para in section_content.split('\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip())
        
        # Save the document
        doc.save(filename)
        
        print(f"Word document lab book saved to {filename}")
        return filename
    
    def add_image_to_document(self, doc_path, image_path, caption=None):
        """Add an image to an existing document"""
        # Check if the document is markdown or docx
        if doc_path.endswith('.md'):
            # Add image to markdown
            with open(doc_path, 'a') as f:
                f.write(f"\n\n![{caption or 'Image'}]({image_path})\n")
                if caption:
                    f.write(f"*{caption}*\n")
            return True
            
        elif doc_path.endswith('.docx'):
            # Add image to docx
            doc = Document(doc_path)
            doc.add_picture(image_path)
            if caption:
                doc.add_paragraph(caption, style='Caption')
            doc.save(doc_path)
            return True
            
        else:
            print(f"Unsupported document format: {doc_path}")
            return False


# Simple command-line test
if __name__ == "__main__":
    generator = DocumentGenerator()
    
    # Test with some sample content
    sample_content = """# Sample Lab Experiment
    
**Date:** 2023-09-20

**Participants:** Dr. Smith, Jane Doe, John Smith

## Objectives

To determine the reaction rate of compound X with catalyst Y.

## Materials and Methods

- Compound X (99% purity)
- Catalyst Y (5 mg)
- 250mL round-bottom flask
- Reflux condenser
- Heating mantle

## Procedure

1. Add 100mL of solvent to the flask
2. Add 5g of compound X
3. Heat to 75°C while stirring
4. Add catalyst Y
5. Maintain temperature for 2 hours

## Results

The reaction yielded 4.2g of product Z (84% yield).
"""
    
    md_file = generator.generate_markdown(sample_content, "Reaction Rate Study")
    docx_file = generator.generate_docx(sample_content, "Reaction Rate Study")
    
    print(f"Test files created: {md_file} and {docx_file}")
